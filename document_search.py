from __future__ import annotations

import re
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Any

import fitz
from docx import Document

from database import get_connection
from search_utils import normalize_standard_no, normalize_text


def _default_docs_dirs() -> list[Path]:
    desktop = Path.home() / "Desktop"
    candidates = [
        desktop / "\u8ba4\u8bc1\u6807\u51c6",
        desktop / "\u4e2d\u6587\u6807\u51c6",
    ]
    return [path for path in candidates if path.exists()]


DEFAULT_DOCS_DIRS = _default_docs_dirs()
DEFAULT_DOCS_DIR = DEFAULT_DOCS_DIRS[0] if DEFAULT_DOCS_DIRS else Path.home() / "Desktop" / "\u8ba4\u8bc1\u6807\u51c6"

QUERY_EXPANSIONS = {
    "\u8776\u9600": ["butterfly valve", "fire-protection butterfly valve", "signal butterfly valve", "supervisory butterfly valve"],
    "butterfly valve": ["\u8776\u9600", "\u6d88\u9632\u8776\u9600", "\u4fe1\u53f7\u8776\u9600"],
    "\u6c34\u6d41\u6307\u793a\u5668": ["waterflow indicator", "water flow switch", "vane type waterflow indicator"],
    "waterflow indicator": ["\u6c34\u6d41\u6307\u793a\u5668", "\u6c34\u6d41\u5f00\u5173"],
    "water flow switch": ["\u6c34\u6d41\u6307\u793a\u5668", "\u6c34\u6d41\u5f00\u5173"],
    "\u6e7f\u5f0f\u62a5\u8b66\u9600": ["alarm valve", "wet alarm valve"],
    "alarm valve": ["\u62a5\u8b66\u9600", "\u6e7f\u5f0f\u62a5\u8b66\u9600", "wet alarm valve"],
    "\u6c34\u529b\u8b66\u94c3": ["water motor gong", "water motor alarm", "mechanical water motor gong"],
    "water motor gong": ["\u6c34\u529b\u8b66\u94c3"],
    "\u538b\u529b\u5f00\u5173": ["pressure switch"],
    "pressure switch": ["\u538b\u529b\u5f00\u5173"],
    "\u5ef6\u8fdf\u5668": ["retard chamber", "retarder"],
    "retard chamber": ["\u5ef6\u8fdf\u5668"],
    "\u95f8\u9600": ["gate valve"],
    "gate valve": ["\u95f8\u9600"],
    "\u6b62\u56de\u9600": ["check valve"],
    "check valve": ["\u6b62\u56de\u9600"],
    "\u55b7\u5934": ["automatic sprinkler", "sprinkler"],
    "automatic sprinkler": ["\u55b7\u5934"],
    "\u6d88\u9632\u6cf5": ["fire pump"],
    "fire pump": ["\u6d88\u9632\u6cf5"],
    "\u6ce1\u6cab": ["foam", "foam concentrate", "foam equipment"],
    "foam": ["\u6ce1\u6cab"],
    "\u6d88\u9632\u8f6f\u7ba1": ["fire hose", "hose"],
    "fire hose": ["\u6d88\u9632\u8f6f\u7ba1"],
    "\u51cf\u538b\u9600": ["pressure reducing valve", "pressure-reducing valve"],
    "pressure reducing valve": ["\u51cf\u538b\u9600"],
    "\u9600\u95e8": ["valve", "gate valve", "check valve", "alarm valve", "butterfly valve"],
    "valve": ["\u9600\u95e8", "\u8776\u9600", "\u95f8\u9600", "\u6b62\u56de\u9600", "\u62a5\u8b66\u9600"],
}


QUICK_SEARCH_TERMS = {
    "\u6e7f\u5f0f\u62a5\u8b66\u9600 / Alarm Valve": "alarm valve",
    "\u8776\u9600 / Butterfly Valve": "butterfly valve",
    "\u6c34\u6d41\u6307\u793a\u5668 / Waterflow Indicator": "waterflow indicator",
    "\u6c34\u529b\u8b66\u94c3 / Water Motor Gong": "water motor gong",
    "\u538b\u529b\u5f00\u5173 / Pressure Switch": "pressure switch",
    "\u5ef6\u8fdf\u5668 / Retard Chamber": "retard chamber",
    "\u95f8\u9600 / Gate Valve": "gate valve",
    "\u6b62\u56de\u9600 / Check Valve": "check valve",
    "\u55b7\u5934 / Automatic Sprinkler": "automatic sprinkler",
    "\u6d88\u9632\u6cf5 / Fire Pump": "fire pump",
    "\u51cf\u538b\u9600 / Pressure Reducing Valve": "pressure reducing valve",
    "UL 193": "UL 193",
    "UL 262": "UL 262",
    "UL 312": "UL 312",
    "UL 753": "UL 753",
    "FM 1041": "FM1041",
    "FM 1055": "1055",
}


COMPONENT_RULES = [
    ("Assembly", "\u603b\u6210/\u9600\u7ec4", ["assembly", "trim", "\u9600\u7ec4", "\u603b\u6210", "\u914d\u7ba1"]),
    ("Valve Body", "\u9600\u4f53", ["body", "valve body", "\u9600\u4f53"]),
    ("Cover", "\u9600\u76d6", ["cover", "handhole cover", "\u9600\u76d6", "\u624b\u5b54\u76d6"]),
    ("Material", "\u6750\u6599", ["material", "cast iron", "bronze", "brass", "rubber", "\u6750\u6599", "\u94f8\u94c1", "\u9752\u94dc", "\u9ec4\u94dc", "\u6a61\u80f6"]),
    ("Flange", "\u6cd5\u5170", ["flange", "flanged", "\u6cd5\u5170"]),
    ("Thread", "\u87ba\u7eb9", ["thread", "threaded", "npt", "\u87ba\u7eb9"]),
    ("Pipe Connection", "\u7ba1\u8def\u8fde\u63a5\u70b9", ["pipe connection", "connection", "inlet", "outlet", "drain", "\u8fde\u63a5", "\u8fdb\u53e3", "\u51fa\u53e3", "\u6392\u6c34"]),
    ("Valve Mechanism", "\u9600\u95e8\u673a\u6784", ["mechanism", "moving parts", "operating parts", "\u673a\u6784", "\u8fd0\u52a8\u90e8\u4ef6"]),
    ("Clapper", "\u9600\u74e3", ["clapper", "disc", "\u9600\u74e3", "\u9600\u76d8"]),
    ("Clapper Support", "\u9600\u74e3\u652f\u6491", ["clapper support", "hinge", "bearing", "pin", "\u9600\u74e3\u652f\u6491", "\u94f0\u94fe", "\u8f74\u627f", "\u9500"]),
    ("Clapper Stop", "\u9600\u74e3\u6321\u5757", ["clapper stop", "stop", "\u9600\u74e3\u6321\u5757", "\u6321\u5757"]),
    ("Seat / Seat Ring", "\u9600\u5ea7/\u5ea7\u5708", ["seat", "seat ring", "\u9600\u5ea7", "\u5ea7\u5708"]),
    ("Retard Chamber", "\u5ef6\u8fdf\u5ba4", ["retard chamber", "retarder", "\u5ef6\u8fdf\u5ba4", "\u5ef6\u8fdf\u5668"]),
    ("Alarm Line", "\u62a5\u8b66\u7ba1\u8def", ["alarm line", "alarm port", "\u62a5\u8b66\u7ba1\u8def", "\u62a5\u8b66\u53e3"]),
    ("Water Motor Gong", "\u6c34\u529b\u8b66\u94c3", ["water motor gong", "water motor alarm", "\u6c34\u529b\u8b66\u94c3"]),
    ("Pressure Switch", "\u538b\u529b\u5f00\u5173", ["pressure switch", "\u538b\u529b\u5f00\u5173"]),
    ("Marking", "\u6807\u8bc6/\u94ed\u724c", ["marking", "mark", "nameplate", "\u6807\u8bc6", "\u94ed\u724c"]),
    ("Installation Instructions", "\u5b89\u88c5\u8bf4\u660e", ["installation", "instructions", "\u5b89\u88c5", "\u8bf4\u660e\u4e66"]),
]


REQUIREMENT_TYPE_RULES = [
    ("Scope", "\u9002\u7528\u8303\u56f4", ["scope", "\u8303\u56f4"]),
    ("Construction", "\u7ed3\u6784\u8981\u6c42", ["construction", "clearance", "arrangement", "\u7ed3\u6784", "\u95f4\u9699"]),
    ("Material", "\u6750\u6599\u8981\u6c42", ["material", "corrosion", "\u6750\u6599", "\u8150\u8680"]),
    ("Dimension", "\u5c3a\u5bf8\u8981\u6c42", ["dimension", "diameter", "size", "inch", "mm", "\u5c3a\u5bf8", "\u76f4\u5f84"]),
    ("Strength Test", "\u5f3a\u5ea6\u8bd5\u9a8c", ["strength", "hydrostatic", "\u5f3a\u5ea6", "\u9759\u6c34\u538b"]),
    ("Leakage Test", "\u5bc6\u5c01/\u6cc4\u6f0f\u8bd5\u9a8c", ["leakage", "leak", "\u6cc4\u6f0f", "\u5bc6\u5c01"]),
    ("Operation Test", "\u52a8\u4f5c/\u64cd\u4f5c\u8bd5\u9a8c", ["operation", "operate", "sensitivity", "\u52a8\u4f5c", "\u64cd\u4f5c", "\u7075\u654f\u5ea6"]),
    ("Durability Test", "\u8010\u4e45\u8bd5\u9a8c", ["durability", "endurance", "\u8010\u4e45"]),
    ("Marking", "\u6807\u8bc6\u8981\u6c42", ["marking", "nameplate", "\u6807\u8bc6", "\u94ed\u724c"]),
    ("Installation", "\u5b89\u88c5\u8bf4\u660e", ["installation", "instructions", "\u5b89\u88c5", "\u8bf4\u660e\u4e66"]),
]


TRANSLATION_GLOSSARY = [
    ("water motor gong", "\u6c34\u529b\u8b66\u94c3"),
    ("water motor alarm", "\u6c34\u529b\u62a5\u8b66\u88c5\u7f6e"),
    ("alarm valve", "\u62a5\u8b66\u9600"),
    ("wet alarm valve", "\u6e7f\u5f0f\u62a5\u8b66\u9600"),
    ("butterfly valve", "\u8776\u9600"),
    ("gate valve", "\u95f8\u9600"),
    ("check valve", "\u6b62\u56de\u9600"),
    ("waterflow indicator", "\u6c34\u6d41\u6307\u793a\u5668"),
    ("pressure switch", "\u538b\u529b\u5f00\u5173"),
    ("retard chamber", "\u5ef6\u8fdf\u5668"),
    ("automatic sprinkler", "\u81ea\u52a8\u55b7\u5934"),
    ("fire pump", "\u6d88\u9632\u6cf5"),
    ("pressure reducing valve", "\u51cf\u538b\u9600"),
    ("hydrostatic strength", "\u9759\u6c34\u538b\u5f3a\u5ea6"),
    ("hydrostatic", "\u9759\u6c34\u538b"),
    ("leakage", "\u6cc4\u6f0f"),
    ("durability", "\u8010\u4e45\u6027"),
    ("performance requirements", "\u6027\u80fd\u8981\u6c42"),
    ("performance", "\u6027\u80fd"),
    ("requirements", "\u8981\u6c42"),
    ("requirement", "\u8981\u6c42"),
    ("test", "\u8bd5\u9a8c"),
    ("inspection", "\u68c0\u9a8c"),
    ("verification", "\u9a8c\u8bc1"),
    ("manufacturer", "\u5236\u9020\u5546"),
    ("installation", "\u5b89\u88c5"),
    ("instructions", "\u8bf4\u660e\u4e66"),
    ("marking", "\u6807\u8bc6"),
    ("sample", "\u6837\u54c1"),
    ("pressure", "\u538b\u529b"),
    ("flow", "\u6d41\u91cf"),
    ("flow rate", "\u6d41\u91cf"),
    ("rated", "\u989d\u5b9a"),
    ("rating", "\u989d\u5b9a\u503c"),
    ("shall", "\u5e94"),
    ("should", "\u5b9c"),
    ("may", "\u53ef"),
    ("without", "\u65e0"),
    ("with", "\u5e26\u6709"),
    ("not less than", "\u4e0d\u5c0f\u4e8e"),
    ("not greater than", "\u4e0d\u5927\u4e8e"),
    ("minimum", "\u6700\u5c0f"),
    ("maximum", "\u6700\u5927"),
    ("inlet", "\u8fdb\u53e3"),
    ("outlet", "\u51fa\u53e3"),
    ("body", "\u672c\u4f53"),
    ("seat", "\u9600\u5ea7"),
    ("gasket", "\u57ab\u7247"),
    ("corrosion", "\u8150\u8680"),
    ("operation", "\u64cd\u4f5c"),
    ("operating", "\u5de5\u4f5c/\u64cd\u4f5c"),
    ("temperature", "\u6e29\u5ea6"),
    ("minutes", "\u5206\u949f"),
    ("minute", "\u5206\u949f"),
    ("hours", "\u5c0f\u65f6"),
    ("hour", "\u5c0f\u65f6"),
    ("psi", "psi"),
    ("gpm", "gpm"),
    ("kpa", "kPa"),
    ("l/min", "L/min"),
]


def init_document_tables() -> None:
    with get_connection() as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS document_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_path TEXT NOT NULL UNIQUE,
                file_name TEXT NOT NULL,
                file_type TEXT,
                source_folder TEXT,
                file_size INTEGER,
                modified_at TEXT,
                indexed_at TEXT DEFAULT (datetime('now','localtime'))
            );

            CREATE TABLE IF NOT EXISTS document_chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_id INTEGER NOT NULL,
                location_label TEXT,
                page_no INTEGER,
                chunk_index INTEGER,
                content TEXT NOT NULL,
                content_norm TEXT NOT NULL,
                FOREIGN KEY(file_id) REFERENCES document_files(id) ON DELETE CASCADE
            );

            CREATE INDEX IF NOT EXISTS idx_document_chunks_file ON document_chunks(file_id);
            CREATE INDEX IF NOT EXISTS idx_document_chunks_page ON document_chunks(page_no);
            """
        )
        existing = {row["name"] for row in conn.execute("PRAGMA table_info(document_files)").fetchall()}
        if "source_folder" not in existing:
            conn.execute("ALTER TABLE document_files ADD COLUMN source_folder TEXT")


def supported_files(folder: Path) -> list[Path]:
    if not folder.exists():
        return []
    return sorted(
        [
            p
            for p in folder.rglob("*")
            if p.is_file() and p.suffix.lower() in {".pdf", ".docx", ".txt", ".md", ".csv"}
        ]
    )


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, max_chars: int = 1600, overlap: int = 0) -> list[str]:
    text = clean_text(text)
    if not text:
        return []
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n|\r?\n", text) if part.strip()]
    if len(paragraphs) > 1:
        chunks: list[str] = []
        current: list[str] = []
        current_len = 0
        for paragraph in paragraphs:
            if len(paragraph) > max_chars:
                if current:
                    chunks.append("\n".join(current).strip())
                    current = []
                    current_len = 0
                chunks.extend(chunk_text(paragraph, max_chars=max_chars, overlap=overlap))
                continue
            if current and current_len + len(paragraph) + 1 > max_chars:
                chunks.append("\n".join(current).strip())
                current = [paragraph]
                current_len = len(paragraph)
            else:
                current.append(paragraph)
                current_len += len(paragraph) + 1
        if current:
            chunks.append("\n".join(current).strip())
        return chunks

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + max_chars)
        part = text[start:end].strip()
        if part:
            chunks.append(part)
        if end >= len(text):
            break
        start = max(0, end - overlap)
    return chunks


def extract_pdf(path: Path) -> list[dict[str, Any]]:
    chunks: list[dict[str, Any]] = []
    with fitz.open(path) as doc:
        for page_index, page in enumerate(doc, start=1):
            text = page.get_text("text")
            for chunk_index, chunk in enumerate(chunk_text(text), start=1):
                chunks.append(
                    {
                        "location_label": f"\u7b2c {page_index} \u9875",
                        "page_no": page_index,
                        "chunk_index": chunk_index,
                        "content": chunk,
                    }
                )
    return chunks


def extract_docx(path: Path) -> list[dict[str, Any]]:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    text = "\n".join(paragraphs)
    return [
        {
            "location_label": f"\u6587\u672c\u7247\u6bb5 {idx}",
            "page_no": None,
            "chunk_index": idx,
            "content": chunk,
        }
        for idx, chunk in enumerate(chunk_text(text), start=1)
    ]


def extract_plain_text(path: Path) -> list[dict[str, Any]]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return [
        {
            "location_label": f"\u6587\u672c\u7247\u6bb5 {idx}",
            "page_no": None,
            "chunk_index": idx,
            "content": chunk,
        }
        for idx, chunk in enumerate(chunk_text(text), start=1)
    ]


def extract_file(path: Path) -> list[dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    return extract_plain_text(path)


def index_documents(folder: Path = DEFAULT_DOCS_DIR) -> dict[str, Any]:
    init_document_tables()
    files = supported_files(folder)
    stats = {"files": 0, "chunks": 0, "errors": []}
    with get_connection() as conn:
        for path in files:
            try:
                chunks = extract_file(path)
                stat = path.stat()
                conn.execute("DELETE FROM document_files WHERE file_path = ?", (str(path),))
                cur = conn.execute(
                    """
                    INSERT INTO document_files
                    (file_path, file_name, file_type, source_folder, file_size, modified_at, indexed_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        str(path),
                        path.name,
                        path.suffix.lower(),
                        str(folder),
                        stat.st_size,
                        datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                        datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    ),
                )
                file_id = int(cur.lastrowid)
                conn.executemany(
                    """
                    INSERT INTO document_chunks
                    (file_id, location_label, page_no, chunk_index, content, content_norm)
                    VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    [
                        (
                            file_id,
                            chunk["location_label"],
                            chunk["page_no"],
                            chunk["chunk_index"],
                            chunk["content"],
                            normalize_text(chunk["content"]),
                        )
                        for chunk in chunks
                    ],
                )
                stats["files"] += 1
                stats["chunks"] += len(chunks)
            except Exception as exc:
                stats["errors"].append(f"{path.name}: {exc}")
    return stats


def index_default_document_folders() -> dict[str, Any]:
    folders = DEFAULT_DOCS_DIRS or [DEFAULT_DOCS_DIR]
    total = {"folders": len(folders), "files": 0, "chunks": 0, "errors": []}
    for folder in folders:
        stats = index_documents(folder)
        total["files"] += stats["files"]
        total["chunks"] += stats["chunks"]
        total["errors"].extend(stats["errors"])
    return total


def expand_query(query: str) -> list[str]:
    q = normalize_text(query)
    terms = [query.strip()]
    for key, values in QUERY_EXPANSIONS.items():
        key_norm = normalize_text(key)
        if key_norm and (key_norm in q or q in key_norm):
            terms.extend(values)

    compact = normalize_standard_no(query)
    if compact.startswith("ul") and len(compact) > 2:
        terms.append(f"UL {compact[2:]}")
        terms.append(compact)

    seen = set()
    unique_terms = []
    for term in terms:
        norm = normalize_text(term)
        if norm and norm not in seen:
            seen.add(norm)
            unique_terms.append(term)
    return unique_terms


def token_terms(query: str) -> list[str]:
    terms = expand_query(query)
    query_norm = normalize_text(query)
    known_cn_terms = [
        "\u6c34\u529b\u8b66\u94c3",
        "\u538b\u529b\u5f00\u5173",
        "\u6c34\u6d41\u6307\u793a\u5668",
        "\u6e7f\u5f0f\u62a5\u8b66\u9600",
        "\u62a5\u8b66\u9600",
        "\u8776\u9600",
        "\u95f8\u9600",
        "\u6b62\u56de\u9600",
        "\u55b7\u5934",
        "\u6d88\u9632\u6cf5",
        "\u68c0\u9a8c",
        "\u68c0\u6d4b",
        "\u8bd5\u9a8c",
        "\u8981\u6c42",
        "\u5224\u5b9a",
    ]
    for term in known_cn_terms:
        if term in query_norm:
            terms.append(term)
    for token in re.findall(r"[\u4e00-\u9fff]{2,}|[a-zA-Z][a-zA-Z0-9/-]{2,}|\d{3,4}", query):
        terms.append(token)
    compact = normalize_standard_no(query)
    standard_match = re.search(r"(ul|fm)(\d{3,4})", compact)
    if standard_match:
        terms.append(standard_match.group(0))
        terms.append(f"{standard_match.group(1).upper()} {standard_match.group(2)}")

    seen = set()
    unique_terms = []
    for term in terms:
        norm = normalize_text(term)
        if norm and norm not in seen:
            seen.add(norm)
            unique_terms.append(term)
    return unique_terms


def intent_terms(query: str) -> list[str]:
    q = normalize_text(query)
    intents = []
    if any(term in q for term in ["\u68c0\u9a8c", "\u68c0\u6d4b", "\u8bd5\u9a8c", "test", "examination"]):
        intents.extend(["\u68c0\u9a8c", "\u68c0\u6d4b", "\u8bd5\u9a8c", "test", "examination", "performance"])
    if any(term in q for term in ["\u8981\u6c42", "requirement"]):
        intents.extend(["\u8981\u6c42", "requirement", "requirements"])
    return intents


def product_terms_in_query(query: str) -> list[str]:
    q = normalize_text(query)
    product_terms = [
        "\u6c34\u529b\u8b66\u94c3",
        "\u538b\u529b\u5f00\u5173",
        "\u6c34\u6d41\u6307\u793a\u5668",
        "\u6e7f\u5f0f\u62a5\u8b66\u9600",
        "\u62a5\u8b66\u9600",
        "\u8776\u9600",
        "\u95f8\u9600",
        "\u6b62\u56de\u9600",
        "\u55b7\u5934",
        "\u6d88\u9632\u6cf5",
        "water motor gong",
        "pressure switch",
        "waterflow indicator",
        "alarm valve",
        "gate valve",
        "check valve",
        "butterfly valve",
    ]
    return [term for term in product_terms if normalize_text(term) in q]


def is_front_matter(row: dict[str, Any]) -> bool:
    content = normalize_text(row.get("content", ""))
    page_no = row.get("page_no")
    if page_no and int(page_no) <= 2:
        front_words = [
            "\u4e2d\u6587\u7ffb\u8bd1\u7a3f",
            "\u539f\u6587\u4ef6",
            "\u53d1\u5e03\u65e5\u671f",
            "\u76ee\u5f55",
            "copyright",
            "translation",
            "cover",
            "contents",
        ]
        return any(word in content for word in front_words)
    return False


def contains_cjk(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text or ""))


def language_of(text: str) -> str:
    return "cn" if contains_cjk(text) else "en"


def offline_translate_en_to_cn(text: str) -> str:
    """Best-effort local reading aid. It preserves numbers and source wording clues."""
    if not text:
        return ""

    normalized = clean_text(text)
    sentences = re.split(r"(?<=[.;:])\s+|\n+", normalized)
    translated_lines: list[str] = []
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        working = sentence
        for english, chinese in sorted(TRANSLATION_GLOSSARY, key=lambda item: len(item[0]), reverse=True):
            working = re.sub(re.escape(english), chinese, working, flags=re.IGNORECASE)
        working = re.sub(r"\bsection\b", "\u7ae0\u8282", working, flags=re.IGNORECASE)
        working = re.sub(r"\bpage\b", "\u9875", working, flags=re.IGNORECASE)
        working = re.sub(r"\bfigure\b", "\u56fe", working, flags=re.IGNORECASE)
        working = re.sub(r"\btable\b", "\u8868", working, flags=re.IGNORECASE)
        translated_lines.append(working)

    return "\n".join(translated_lines[:8])


def extract_doc_key(text: str) -> str:
    compact = normalize_standard_no(text)
    match = re.search(r"(ul|fm)(\d{3,4})", compact)
    if match:
        return match.group(0)
    match = re.search(r"\b(1055|1041)\b", normalize_text(text))
    return match.group(1) if match else ""


def make_snippet(content: str, terms: list[str], width: int = 260) -> str:
    norm_content = normalize_text(content)
    positions = []
    for term in terms:
        norm_term = normalize_text(term)
        if norm_term:
            pos = norm_content.find(norm_term)
            if pos >= 0:
                positions.append(pos)
    if not positions:
        return content[:width].strip()
    center = min(positions)
    start = max(0, center - width // 2)
    end = min(len(content), start + width)
    snippet = content[start:end].strip()
    if start > 0:
        snippet = "..." + snippet
    if end < len(content):
        snippet += "..."
    return snippet


def search_documents(query: str, file_type: str = "", limit: int = 80) -> list[dict[str, Any]]:
    init_document_tables()
    terms = token_terms(query)
    if not terms:
        return []

    rows = []
    with get_connection() as conn:
        conn.row_factory = sqlite3.Row
        all_rows = conn.execute(
            """
            SELECT
                c.id AS chunk_id,
                c.location_label,
                c.page_no,
                c.chunk_index,
                c.content,
                c.content_norm,
                f.file_name,
                f.file_path,
                f.file_type,
                f.source_folder,
                f.modified_at,
                f.indexed_at
            FROM document_chunks c
            JOIN document_files f ON f.id = c.file_id
            WHERE (? = '' OR f.file_type = ?)
            """,
            (file_type, file_type),
        ).fetchall()

    non_standard_terms = [
        term
        for term in terms
        if not re.fullmatch(r"(ul|fm)\s*\d{3,4}|\d{3,4}", normalize_text(term), flags=re.IGNORECASE)
    ]
    query_intents = intent_terms(query)
    query_products = product_terms_in_query(query)

    for row in all_rows:
        row_dict = dict(row)
        content_norm = row_dict["content_norm"]
        content_score = 0
        file_score = 0
        matched_non_standard_count = 0
        matched_intent_count = 0
        non_standard_content_hit = False
        for term in terms:
            norm = normalize_text(term)
            compact = normalize_standard_no(term)
            if norm and norm in content_norm:
                weight = 20 if " " in norm or contains_cjk(norm) else 8
                content_score += weight + min(content_norm.count(norm), 5) * 3
                if term in non_standard_terms:
                    non_standard_content_hit = True
                    matched_non_standard_count += 1
            if compact and compact in normalize_standard_no(row_dict["file_name"]):
                file_score += 30
            if norm and norm in normalize_text(row_dict["file_name"]):
                file_score += 25

        for intent in query_intents:
            if normalize_text(intent) in content_norm:
                matched_intent_count += 1
        if query_intents:
            content_score += matched_intent_count * 18
        for product_term in query_products:
            product_norm = normalize_text(product_term)
            if product_norm in content_norm:
                content_score += min(content_norm.count(product_norm), 4) * 20

        if non_standard_terms and not non_standard_content_hit:
            continue

        coverage_bonus = matched_non_standard_count * 15
        clause_bonus = 0
        if re.search(r"(^|\s)\d+(\.\s*\d+)+", row_dict.get("content", "")):
            clause_bonus += 25
        if re.search(r"\d+(\.\d+)?\s*(mpa|psi|kpa|gpm|l/min|db|℃|s|min|\u79d2|\u5206\u949f)", content_norm, re.IGNORECASE):
            clause_bonus += 20
        if query_intents and any(
            term in content_norm
            for term in [
                "\u68c0\u6d4b/\u9a8c\u8bc1",
                "\u6d4b\u8bd5/\u9a8c\u8bc1",
                "\u6027\u80fd\u8981\u6c42",
                "\u8bd5\u9a8c\u65b9\u6cd5",
                "\u53ef\u542c\u6027",
                "\u8010\u4e45\u6027",
                "performance requirements",
                "test method",
                "test/verification",
            ]
        ):
            clause_bonus += 80
        front_penalty = 110 if is_front_matter(row_dict) else 0
        admin_penalty = 0
        if query_intents and any(
            term in content_norm
            for term in [
                "\u5236\u9020\u5546\u8981\u6c42",
                "\u5236\u9020\u5546\u9700",
                "\u8d28\u91cf\u63a7\u5236",
                "\u8d28\u91cf\u4fdd\u8bc1\u4f53\u7cfb",
                "\u76d1\u7763\u5ba1\u6838",
                "\u76d1\u7763\u8ba1\u5212",
                "\u8ba4\u8bc1\u673a\u6784",
                "\u8ba4\u8bc1\u62a5\u544a",
                "\u56fe\u7eb8\u548c\u53d8\u66f4",
                "manufacturer requirements",
                "quality control",
                "quality assurance",
                "surveillance audit",
            ]
        ):
            admin_penalty += 120
        if query_intents and any(
            term in content_norm
            for term in [
                "\u5355\u4f4d\u5236",
                "\u89c4\u8303\u6027\u5f15\u7528\u6587\u4ef6",
                "\u672f\u8bed\u548c\u5b9a\u4e49",
                "\u9002\u7528\u8303\u56f4",
                "units of measurement",
                "referenced publications",
                "terms and definitions",
            ]
        ):
            admin_penalty += 120
        if "\u76ee\u5f55" in content_norm or "................" in row_dict.get("content", ""):
            admin_penalty += 65
        score = content_score + min(file_score, 40) + coverage_bonus + clause_bonus - front_penalty - admin_penalty
        if score > 0:
            row_dict["score"] = score
            row_dict["matched_terms"] = ", ".join(terms)
            row_dict["snippet"] = make_snippet(row_dict["content"], terms, width=700)
            row_dict["language"] = language_of(row_dict["content"])
            rows.append(row_dict)

    rows.sort(key=lambda item: (item["score"], item["file_name"], -(item["page_no"] or 0)), reverse=True)
    return rows[:limit]


def find_bilingual_counterpart(query: str, source_row: dict[str, Any]) -> dict[str, Any] | None:
    source_lang = source_row.get("language") or language_of(source_row.get("content", ""))
    target_lang = "en" if source_lang == "cn" else "cn"
    source_key = extract_doc_key(source_row.get("file_name", "") + " " + source_row.get("file_path", ""))
    candidates = search_documents(query, limit=200)

    ranked = []
    same_standard_ranked = []
    for candidate in candidates:
        if candidate["chunk_id"] == source_row.get("chunk_id"):
            continue
        if candidate.get("language") != target_lang:
            continue
        candidate_key = extract_doc_key(candidate.get("file_name", "") + " " + candidate.get("file_path", ""))
        bonus = 0
        if source_key and candidate_key and source_key == candidate_key:
            bonus += 80
        elif source_key:
            continue
        if source_row.get("page_no") and candidate.get("page_no"):
            distance = abs(int(source_row["page_no"]) - int(candidate["page_no"]))
            bonus += max(0, 25 - distance)
        target = (candidate["score"] + bonus, candidate)
        ranked.append(target)
        if source_key and candidate_key and source_key == candidate_key:
            same_standard_ranked.append(target)

    if same_standard_ranked:
        same_standard_ranked.sort(key=lambda item: item[0], reverse=True)
        result = same_standard_ranked[0][1]
        result["pair_score"] = same_standard_ranked[0][0]
        return result
    if not ranked:
        return None
    ranked.sort(key=lambda item: item[0], reverse=True)
    result = ranked[0][1]
    result["pair_score"] = ranked[0][0]
    return result


def _match_rules(content: str, rules: list[tuple[str, str, list[str]]]) -> list[tuple[str, str]]:
    norm = normalize_text(content)
    hits: list[tuple[str, str]] = []
    for en, cn, keywords in rules:
        if any(normalize_text(keyword) in norm for keyword in keywords):
            hits.append((en, cn))
    return hits


def list_indexed_files() -> list[dict[str, Any]]:
    init_document_tables()
    with get_connection() as conn:
        rows = conn.execute(
            """
            SELECT id, file_name, file_path, file_type, source_folder, modified_at, indexed_at
            FROM document_files
            ORDER BY file_name
            """
        ).fetchall()
    return [dict(row) for row in rows]


def split_file_candidates(file_id: int | None = None, limit: int = 300) -> list[dict[str, Any]]:
    """Classify indexed chunks into component-level requirement candidates."""
    init_document_tables()
    where = "WHERE f.id = ?" if file_id else ""
    params = (file_id,) if file_id else ()
    with get_connection() as conn:
        rows = conn.execute(
            f"""
            SELECT
                c.id AS chunk_id,
                c.location_label,
                c.page_no,
                c.chunk_index,
                c.content,
                f.id AS file_id,
                f.file_name,
                f.file_path,
                f.source_folder,
                f.file_type
            FROM document_chunks c
            JOIN document_files f ON f.id = c.file_id
            {where}
            ORDER BY f.file_name, COALESCE(c.page_no, 0), c.chunk_index
            """,
            params,
        ).fetchall()

    candidates: list[dict[str, Any]] = []
    for row in rows:
        item = dict(row)
        component_hits = _match_rules(item["content"], COMPONENT_RULES)
        type_hits = _match_rules(item["content"], REQUIREMENT_TYPE_RULES)
        if not component_hits and not type_hits:
            continue
        if not component_hits:
            component_hits = [("General", "\u901a\u7528\u8981\u6c42")]
        if not type_hits:
            type_hits = [("General", "\u672a\u5206\u7c7b")]

        component_en, component_cn = component_hits[0]
        requirement_type_en, requirement_type_cn = type_hits[0]
        candidates.append(
            {
                "file_id": item["file_id"],
                "file_name": item["file_name"],
                "file_path": item["file_path"],
                "source_folder": item.get("source_folder"),
                "chunk_id": item["chunk_id"],
                "location_label": item["location_label"],
                "page_no": item["page_no"],
                "content": item["content"],
                "language": language_of(item["content"]),
                "score": 0,
                "component_cn": component_cn,
                "component_en": component_en,
                "requirement_type": requirement_type_en,
                "requirement_type_cn": requirement_type_cn,
                "requirement_object": f"{component_en} - {requirement_type_en}",
                "candidate_snippet": item["content"],
                "source_excerpt_short": item["content"][:220],
            }
        )
        if len(candidates) >= limit:
            break
    return candidates


def split_standard_candidates(standard_no: str, limit: int = 300) -> list[dict[str, Any]]:
    """Return component candidates from files whose names match a standard number."""
    standard_key = normalize_standard_no(standard_no)
    if not standard_key:
        return []

    files = list_indexed_files()
    matched_file_ids = [
        item["id"]
        for item in files
        if standard_key in normalize_standard_no(item.get("file_name", ""))
        or standard_key in normalize_standard_no(item.get("file_path", ""))
    ]
    candidates: list[dict[str, Any]] = []
    seen = set()
    for file_id in matched_file_ids:
        for item in split_file_candidates(file_id, limit=limit):
            key = (item["file_name"], item["location_label"], item["component_en"], item["requirement_type"])
            if key in seen:
                continue
            seen.add(key)
            candidates.append(item)
            if len(candidates) >= limit:
                return candidates
    return candidates


def document_index_summary() -> dict[str, Any]:
    init_document_tables()
    with get_connection() as conn:
        files = conn.execute("SELECT COUNT(*) FROM document_files").fetchone()[0]
        chunks = conn.execute("SELECT COUNT(*) FROM document_chunks").fetchone()[0]
        last_indexed = conn.execute("SELECT MAX(indexed_at) FROM document_files").fetchone()[0]
    return {"files": files, "chunks": chunks, "last_indexed": last_indexed or "\u5c1a\u672a\u7d22\u5f15"}
